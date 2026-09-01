import streamlit as st
import pandas as pd
import datetime
import os
import hashlib
import smtplib
from io import BytesIO
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# Librerías para generación de PDF y Word
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# ---------------------------------------------------------
# CONFIGURACIÓN DE PÁGINA
# ---------------------------------------------------------
st.set_page_config(page_title="Control Financiero CGPA", layout="wide")

ARCHIVO_DATOS = "datos_cgpa.csv"
ARCHIVO_USUARIOS = "usuarios.csv"
CARPETA_COMPROBANTES = "comprobantes"
LIMITE_MB = 20
TAMANO_MAXIMO_BYTES = LIMITE_MB * 1024 * 1024

if not os.path.exists(CARPETA_COMPROBANTES):
    os.makedirs(CARPETA_COMPROBANTES)

# ---------------------------------------------------------
# GESTIÓN DE USUARIOS
# ---------------------------------------------------------
def cargar_usuarios():
    if os.path.exists(ARCHIVO_USUARIOS):
        return pd.read_csv(ARCHIVO_USUARIOS)
    else:
        admin_defecto = pd.DataFrame([{
            "Usuario": "admin",
            "Password_Hash": hashlib.sha256("admin123".encode()).hexdigest(),
            "Nombre": "Administrador General",
            "Rol": "Administrador",
            "Estado_Aprobacion": "APROBADO",
            "Tipo_Login": "Local"
        }])
        admin_defecto.to_csv(ARCHIVO_USUARIOS, index=False)
        return admin_defecto

def guardar_usuarios(df_usr):
    df_usr.to_csv(ARCHIVO_USUARIOS, index=False)

def hash_pass(password):
    return hashlib.sha256(password.encode()).hexdigest()

# ---------------------------------------------------------
# GESTIÓN DE DATOS Y BALANCE
# ---------------------------------------------------------
def cargar_datos():
    if os.path.exists(ARCHIVO_DATOS):
        try:
            df = pd.read_csv(ARCHIVO_DATOS)
            df["Monto_Original_CLP"] = pd.to_numeric(df["Monto_Original_CLP"], errors="coerce").fillna(0)
            df["Monto_Corregido_CLP"] = pd.to_numeric(df["Monto_Corregido_CLP"], errors="coerce").fillna(0)
            df["Fecha_Transaccion"] = pd.to_datetime(df["Fecha_Transaccion"], errors="coerce")
            return df
        except Exception:
            return crear_dataframe_vacio()
    else:
        return crear_dataframe_vacio()

def crear_dataframe_vacio():
    return pd.DataFrame(columns=[
        "ID", "Fecha_Transaccion", "Tipo_Movimiento", "Categoria", 
        "Detalle", "Responsable", "Medio_Pago", "Monto_Original_CLP", 
        "Monto_Corregido_CLP", "Usuario_Corrigio", "Ruta_Comprobante", 
        "Usuario_Registro", "Estado_Registro"
    ])

def guardar_datos(df):
    df.to_csv(ARCHIVO_DATOS, index=False)

def calcular_balance(df):
    if df.empty or "Estado_Registro" not in df.columns:
        return 0, 0, 0
    df_activos = df[df["Estado_Registro"] == "ACTIVO"].copy()
    if df_activos.empty:
        return 0, 0, 0
    df_activos["Monto_Efectivo"] = df_activos.apply(
        lambda r: r["Monto_Corregido_CLP"] if r["Monto_Corregido_CLP"] > 0 else r["Monto_Original_CLP"], axis=1
    )
    ingresos = df_activos[df_activos["Tipo_Movimiento"] == "Ingreso"]["Monto_Efectivo"].sum()
    egresos = df_activos[df_activos["Tipo_Movimiento"] == "Egreso"]["Monto_Efectivo"].sum()
    return ingresos, egresos, ingresos - egresos

# ---------------------------------------------------------
# GENERACIÓN DE REPORTES (EXCEL, PDF, WORD)
# ---------------------------------------------------------
def generar_excel(df):
    buffer = BytesIO()
    df_export = df.copy()
    if "Fecha_Transaccion" in df_export.columns:
        df_export["Fecha_Transaccion"] = df_export["Fecha_Transaccion"].dt.strftime('%Y-%m-%d')
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df_export.to_excel(writer, index=False, sheet_name="Reporte_Financiero")
    buffer.seek(0)
    return buffer.getvalue()

def generar_pdf(df, ingresos, egresos, saldo, titulo_periodo="GENERAL"):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, 750, f"REPORTE FINANCIERO OFICIAL CGPA - PERÍODO: {titulo_periodo}")
    p.setFont("Helvetica", 10)
    p.drawString(50, 735, f"Fecha de Emisión: {datetime.date.today()} | Documento Inalterable")
    p.line(50, 725, 550, 725)
    
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, 695, f"Ingresos Totales: $ {ingresos:,.0f} CLP")
    p.drawString(50, 675, f"Egresos Totales: $ {egresos:,.0f} CLP")
    p.drawString(50, 655, f"Balance Disponible: $ {saldo:,.0f} CLP")
    p.line(50, 640, 550, 640)
    
    y = 610
    p.setFont("Helvetica-Bold", 10)
    p.drawString(50, y, "ID | Fecha | Tipo | Categoría | Monto Efectivo CLP")
    p.setFont("Helvetica", 9)
    
    for idx, row in df.iterrows():
        y -= 20
        if y < 50:
            p.showPage()
            y = 750
            p.setFont("Helvetica", 9)
        monto = row["Monto_Corregido_CLP"] if row["Monto_Corregido_CLP"] > 0 else row["Monto_Original_CLP"]
        fecha_str = row['Fecha_Transaccion'].strftime('%Y-%m-%d') if pd.notnull(row['Fecha_Transaccion']) else "N/A"
        p.drawString(50, y, f"#{row['ID']} | {fecha_str} | {row['Tipo_Movimiento']} | {row['Categoria']} | $ {monto:,.0f}")
        
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer.getvalue()

def generar_word(df, ingresos, egresos, saldo, titulo_periodo="GENERAL"):
    doc = Document()
    doc.add_heading(f"Reporte Financiero Oficial CGP - {titulo_periodo}", 0)
    doc.add_paragraph(f"Fecha de generación: {datetime.date.today()}")
    
    doc.add_heading("Resumen del Balance", level=1)
    doc.add_paragraph(f"• Ingresos Totales: $ {ingresos:,.0f} CLP")
    doc.add_paragraph(f"• Egresos Totales: $ {egresos:,.0f} CLP")
    doc.add_paragraph(f"• Balance Disponible: $ {saldo:,.0f} CLP")
    
    doc.add_heading("Detalle de Transacciones", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Fecha'
    hdr_cells[2].text = 'Tipo'
    hdr_cells[3].text = 'Categoría'
    hdr_cells[4].text = 'Monto CLP'
    
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        monto = row["Monto_Corregido_CLP"] if row["Monto_Corregido_CLP"] > 0 else row["Monto_Original_CLP"]
        fecha_str = row['Fecha_Transaccion'].strftime('%Y-%m-%d') if pd.notnull(row['Fecha_Transaccion']) else "N/A"
        row_cells[0].text = str(row['ID'])
        row_cells[1].text = fecha_str
        row_cells[2].text = str(row['Tipo_Movimiento'])
        row_cells[3].text = str(row['Categoria'])
        row_cells[4].text = f"$ {monto:,.0f}"
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def enviar_correo(destinatario, asunto, mensaje, archivo_bytes, nombre_archivo):
    remitente = "tu_correo@gmail.com"
    password = "tu_password_de_aplicacion"
    
    msg = MIMEMultipart()
    msg['From'] = remitente
    msg['To'] = destinatario
    msg['Subject'] = asunto
    msg.attach(MIMEText(mensaje, 'plain'))
    
    adjunto = MIMEApplication(archivo_bytes, Name=nombre_archivo)
    adjunto['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
    msg.attach(adjunto)
    
    try:
        if remitente == "tu_correo@gmail.com":
            return True, "Simulación de envío exitosa: Correo preparado correctamente. (Configura SMTP para envío real)."
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(remitente, password)
        server.send_message(msg)
        server.quit()
        return True, "Correo enviado exitosamente."
    except Exception as e:
        return False, f"Error al enviar el correo: {e}"

# ---------------------------------------------------------
# AUTENTICACIÓN Y LOGIN
# ---------------------------------------------------------
if "usuario_logueado" not in st.session_state:
    st.session_state.usuario_logueado = None

df_usuarios = cargar_usuarios()

if st.session_state.usuario_logueado is None:
    st.title("🔑 Control de Acceso y Autenticación - CGP")
    
    tab_login, tab_registro = st.tabs(["Iniciar Sesión", "Registrarse"])
    
    with tab_login:
        usr_input = st.text_input("Usuario:")
        pass_input = st.text_input("Contraseña:", type="password")
        if st.button("Ingresar"):
            usr_match = df_usuarios[(df_usuarios["Usuario"] == usr_input) & (df_usuarios["Password_Hash"] == hash_pass(pass_input))]
            if not usr_match.empty:
                estado = usr_match.iloc[0]["Estado_Aprobacion"]
                if estado == "APROBADO":
                    st.session_state.usuario_logueado = usr_match.iloc[0].to_dict()
                    st.success("✅ Acceso autorizado.")
                    st.rerun()
                else:
                    st.warning("⏳ Tu cuenta requiere aprobación del Administrador.")
            else:
                st.error("❌ Usuario o contraseña incorrectos.")
                
    with tab_registro:
        nuevo_usr = st.text_input("Nuevo Usuario:")
        nuevo_nom = st.text_input("Nombre Completo:")
        nuevo_pass = st.text_input("Nueva Contraseña:", type="password")
        if st.button("Solicitar Registro"):
            if nuevo_usr and nuevo_pass:
                nueva_cuenta = {
                    "Usuario": nuevo_usr,
                    "Password_Hash": hash_pass(nuevo_pass),
                    "Nombre": nuevo_nom,
                    "Rol": "Tesorero/Operador",
                    "Estado_Aprobacion": "PENDIENTE",
                    "Tipo_Login": "Local"
                }
                df_usuarios = pd.concat([df_usuarios, pd.DataFrame([nueva_cuenta])], ignore_index=True)
                guardar_usuarios(df_usuarios)
                st.info("📩 Registro solicitado. Esperando validación del Administrador.")

    st.stop()

# ---------------------------------------------------------
# APLICACIÓN PRINCIPAL
# ---------------------------------------------------------
usuario_actual = st.session_state.usuario_logueado
es_administrador = (usuario_actual["Rol"] == "Administrador")

st.sidebar.write(f"👤 Usuario: **{usuario_actual['Nombre']}**")
st.sidebar.write(f"🔑 Rol: **{usuario_actual['Rol']}**")
if st.sidebar.button("Cerrar Sesión"):
    st.session_state.usuario_logueado = None
    st.rerun()

st.title("🛡️ Sistema Financiero y Balance Contable - CGP")

df_transacciones = cargar_datos()
ing_act, egr_act, saldo_act = calcular_balance(df_transacciones)

# Métrica de balance general
st.markdown("### 💳 Balance Actual de la Cuenta")
col1, col2, col3 = st.columns(3)
col1.metric("Ingresos Acumulados", f"$ {ing_act:,.0f} CLP")
col2.metric("Egresos Acumulados", f"$ {egr_act:,.0f} CLP")
col3.metric("BALANCE DISPONIBLE", f"$ {saldo_act:,.0f} CLP")

st.divider()

# Pestañas principales
tab_nuevo, tab_corregir, tab_historial, tab_exportar = st.tabs([
    "📥 Nuevo Movimiento", 
    "✏️ Corregir Monto", 
    "📊 Historial", 
    "📥 Descargar y Enviar Reporte"
])

# PESTAÑA 1: NUEVO MOVIMIENTO
with tab_nuevo:
    with st.form("form_registro", clear_on_submit=True):
        col_a, col_b = st.columns(2)
        with col_a:
            tipo = st.selectbox("Tipo de Movimiento:", ["Ingreso", "Egreso"])
            monto = st.number_input("Monto ($ CLP):", min_value=1, step=1000)
            responsable = st.text_input("Responsable / Destinatario:")
            categoria = st.selectbox("Categoría:", ["Fondo Catastrófico", "Cuota Incorporación", "Gastos Operativos", "Otros"])
        
        with col_b:
            medio_pago = st.selectbox("Medio de Pago:", ["Transferencia / Cuenta", "Efectivo", "Cheque"])
            detalle = st.text_area("Detalle / Observaciones:")
            imagen_adjunta = st.file_uploader(
                "🔴 Comprobante (OBLIGATORIO - Máx 20MB):", 
                type=["jpg", "jpeg", "png", "webp", "bmp"]
            )

        btn_guardar = st.form_submit_button("💾 Guardar y Actualizar Balance")

        if btn_guardar:
            if imagen_adjunta is None:
                st.error("❌ Es obligatorio adjuntar el comprobante.")
            elif imagen_adjunta.size > TAMANO_MAXIMO_BYTES:
                st.error(f"❌ El archivo excede el límite de {LIMITE_MB} MB.")
            else:
                nuevo_id = len(df_transacciones) + 1
                extension = imagen_adjunta.name.split(".")[-1]
                nombre_archivo = f"comprobante_{nuevo_id}.{extension}"
                ruta_destino = os.path.join(CARPETA_COMPROBANTES, nombre_archivo)

                with open(ruta_destino, "wb") as f:
                    f.write(imagen_adjunta.getbuffer())

                nueva_fila = {
                    "ID": nuevo_id,
                    "Fecha_Transaccion": str(datetime.date.today()),
                    "Tipo_Movimiento": tipo,
                    "Categoria": categoria,
                    "Detalle": detalle,
                    "Responsable": responsable,
                    "Medio_Pago": medio_pago,
                    "Monto_Original_CLP": monto,
                    "Monto_Corregido_CLP": 0,
                    "Usuario_Corrigio": "Sin Corrección",
                    "Ruta_Comprobante": ruta_destino,
                    "Usuario_Registro": usuario_actual["Usuario"],
                    "Estado_Registro": "ACTIVO"
                }

                df_transacciones = pd.concat([df_transacciones, pd.DataFrame([nueva_fila])], ignore_index=True)
                guardar_datos(df_transacciones)
                st.success(f"✅ Transacción #{nuevo_id} guardada con éxito.")
                st.rerun()

# PESTAÑA 2: CORREGIR MONTO
with tab_corregir:
    df_activos = df_transacciones[df_transacciones["Estado_Registro"] == "ACTIVO"]
    if not df_activos.empty:
        id_selec = st.selectbox("Selecciona ID a Corregir:", df_activos["ID"].tolist())
        reg_actual = df_activos[df_activos["ID"] == id_selec].iloc[0]
        monto_mostrar = reg_actual["Monto_Corregido_CLP"] if reg_actual["Monto_Corregido_CLP"] > 0 else reg_actual["Monto_Original_CLP"]
        
        with st.form("form_corregir"):
            nuevo_monto_input = st.number_input("Nuevo Monto Correcto ($ CLP):", min_value=1, step=1000, value=int(monto_mostrar))
            btn_aplicar = st.form_submit_button("💾 Guardar Corrección")
            if btn_aplicar:
                idx = df_transacciones.index[df_transacciones["ID"] == id_selec].tolist()[0]
                df_transacciones.at[idx, "Monto_Corregido_CLP"] = nuevo_monto_input
                df_transacciones.at[idx, "Usuario_Corrigio"] = usuario_actual["Usuario"]
                guardar_datos(df_transacciones)
                st.success("✅ Monto corregido con éxito.")
                st.rerun()

# PESTAÑA 3: HISTORIAL
with tab_historial:
    if not df_transacciones.empty:
        df_mostrar = df_transacciones[df_transacciones["Estado_Registro"] == "ACTIVO"].copy()
        if "Fecha_Transaccion" in df_mostrar.columns:
            df_mostrar["Fecha_Transaccion"] = df_mostrar["Fecha_Transaccion"].dt.strftime('%Y-%m-%d')
        st.dataframe(df_mostrar, use_container_width=True)

# PESTAÑA 4: DESCARGAR Y ENVIAR REPORTE
with tab_exportar:
    st.subheader("📥 Generación de Reportes y Envíos por Correo")
    
    df_activos = df_transacciones[df_transacciones["Estado_Registro"] == "ACTIVO"].copy() if not df_transacciones.empty else crear_dataframe_vacio()
    
    # FILTRO POR PERÍODO Y SEMESTRE
    st.markdown("#### 🗓️ Selección de Período de Reporte")
    col_p1, col_p2 = st.columns(2)
    
    # Obtener años disponibles
    if not df_activos.empty and "Fecha_Transaccion" in df_activos.columns:
        anos_disponibles = sorted(df_activos["Fecha_Transaccion"].dt.year.dropna().unique().astype(int).tolist())
        if not anos_disponibles:
            anos_disponibles = [datetime.date.today().year]
    else:
        anos_disponibles = [datetime.date.today().year]
        
    with col_p1:
        opciones_ano = ["Todos los Años"] + [str(a) for a in anos_disponibles]
        ano_seleccionado = st.selectbox("Seleccionar Año:", opciones_ano)
        
    with col_p2:
        semestre_seleccionado = st.selectbox("Seleccionar Semestre:", ["Todos los Semestres", "Semestre 1 (Ene - Jun)", "Semestre 2 (Jul - Dic)"])

    # Filtrar el DataFrame según la selección del usuario
    df_filtrado = df_activos.copy()
    
    if ano_seleccionado != "Todos los Años" and not df_filtrado.empty:
        df_filtrado = df_filtrado[df_filtrado["Fecha_Transaccion"].dt.year == int(ano_seleccionado)]
        
    if semestre_seleccionado == "Semestre 1 (Ene - Jun)" and not df_filtrado.empty:
        df_filtrado = df_filtrado[df_filtrado["Fecha_Transaccion"].dt.month.isin([1, 2, 3, 4, 5, 6])]
    elif semestre_seleccionado == "Semestre 2 (Jul - Dic)" and not df_filtrado.empty:
        df_filtrado = df_filtrado[df_filtrado["Fecha_Transaccion"].dt.month.isin([7, 8, 9, 10, 11, 12])]

    # Recalcular saldos para el reporte filtrado
    ing_filt, egr_filt, saldo_filt = calcular_balance(df_filtrado)
    titulo_periodo = f"{ano_seleccionado} - {semestre_seleccionado}"

    st.info(f"📊 **Período Seleccionado:** {titulo_periodo} | **Registros:** {len(df_filtrado)} | **Saldo Período:** $ {saldo_filt:,.0f} CLP")

    st.divider()

    # DESCARGA SEGÚN EL ROL DE USUARIO
    st.markdown("#### 📄 Opciones de Descarga Directa")
    
    pdf_bytes = generar_pdf(df_filtrado, ing_filt, egr_filt, saldo_filt, titulo_periodo)
    
    if es_administrador:
        st.success("🔓 **Perfil Administrador:** Tienes habilitadas todas las opciones de descarga (PDF, Excel editable y Word).")
        col_d1, col_d2, col_d3 = st.columns(3)
        
        with col_d1:
            st.download_button(
                label="📄 Descargar PDF (.pdf)",
                data=pdf_bytes,
                file_name=f"Reporte_CGPA_{titulo_periodo}.pdf",
                mime="application/pdf"
            )
        with col_d2:
            excel_bytes = generar_excel(df_filtrado)
            st.download_button(
                label="📊 Descargar Excel (.xlsx)",
                data=excel_bytes,
                file_name=f"Reporte_CGPA_{titulo_periodo}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        with col_d3:
            word_bytes = generar_word(df_filtrado, ing_filt, egr_filt, saldo_filt, titulo_periodo)
            st.download_button(
                label="📝 Descargar Word (.docx)",
                data=word_bytes,
                file_name=f"Reporte_CGPA_{titulo_periodo}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("🔒 **Perfil Usuario / Tesorero:** Por normas de seguridad e inviolabilidad contable, únicamente tienes permitido descargar archivos oficiales en **formato PDF**.")
        st.download_button(
            label="📄 Descargar Reporte Oficial Inalterable (.pdf)",
            data=pdf_bytes,
            file_name=f"Reporte_Oficial_CGPA_{titulo_periodo}.pdf",
            mime="application/pdf"
        )

    st.divider()

    # ENVÍO DE CORREO
    st.subheader("✉️ Redirigir o Enviar Reporte por Correo Electrónico")
    
    with st.form("form_email"):
        correo_destino = st.text_input("Correo Electrónico Destinatario:")
        asunto_email = st.text_input("Asunto:", value=f"Reporte Financiero CGPA - {titulo_periodo}")
        
        if es_administrador:
            formato_adjunto = st.selectbox("Seleccionar Formato del Archivo Adjunto:", ["PDF (.pdf)", "Excel (.xlsx)", "Word (.docx)"])
        else:
            st.info("🔒 Para usuarios no administradores, el reporte se adjuntará exclusivamente en formato PDF.")
            formato_adjunto = "PDF (.pdf)"
            
        mensaje_body = st.text_area("Mensaje:", value=f"Estimados,\n\nAdjunto el reporte de ingresos y egresos correspondiente al período {titulo_periodo}.\n\nSaludos cordiales.")
        
        btn_enviar_email = st.form_submit_button("📤 Enviar Correo con Adjunto")
        
        if btn_enviar_email:
            if not correo_destino.strip():
                st.error("❌ Debe ingresar una dirección de correo válida.")
            else:
                if "Excel" in formato_adjunto and es_administrador:
                    bytes_enviar = generar_excel(df_filtrado)
                    nombre_adj = f"Reporte_CGPA_{titulo_periodo}.xlsx"
                elif "Word" in formato_adjunto and es_administrador:
                    bytes_enviar = generar_word(df_filtrado, ing_filt, egr_filt, saldo_filt, titulo_periodo)
                    nombre_adj = f"Reporte_CGPA_{titulo_periodo}.docx"
                else:
                    bytes_enviar = pdf_bytes
                    nombre_adj = f"Reporte_CGPA_{titulo_periodo}.pdf"
                
                exito, msg_respuesta = enviar_correo(correo_destino, asunto_email, mensaje_body, bytes_enviar, nombre_adj)
                if exito:
                    st.success(f"✅ {msg_respuesta}")
                else:
                    st.error(f"❌ {msg_respuesta}")